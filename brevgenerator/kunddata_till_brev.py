import os
import sys
import csv
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Inches

# Get the directory of the current script
script_dir = os.path.dirname(os.path.realpath(__file__))


def get_resource_path(relative_path):
    """Constructs a path to the resource that works whether the script is run from source or from a bundled executable."""
    if getattr(sys, 'frozen', False):
        # If the application is run as a bundled executable, determine the directory
        # containing the executable and construct a path relative to it.
        base_path = os.path.dirname(sys.executable)
    else:
        # If the application is run in a development environment, use a directory
        # relative to the script file. This might be adjusted based on your project structure.
        base_path = os.path.dirname(os.path.abspath(__file__))

    # Construct the full path to the resource.
    resource_path = os.path.join(base_path, relative_path)
    return resource_path

# Example usage:
# Assuming 'nv4.docx' and 'kunddata.csv' are in a directory named 'resources' located
# in the same directory as the executable or script.
nv4_docx_path = get_resource_path('nv4.docx')
kunddata_csv_path = get_resource_path('kunddata.csv')

# Use the get_resource_path function to find nv4.docx and kunddata.csv
kunddata = get_resource_path('kunddata.csv')
template = get_resource_path('nv4.docx')
today = datetime.date.today().strftime("%Y-%m-%d") # Correct way to get today's date
output_path = os.path.join(os.path.dirname(template), today + '.docx')

class Kund:
    def __init__(self, kundnr, fornamn, efternamn, adress1, adress21, adress22, email, phone):
        self.kundnr = kundnr
        self.fornamn = fornamn
        self.efternamn = efternamn
        self.adress1 = adress1
        self.adress21 = adress21
        self.adress22 = adress22
        self.email = email
        self.phone = phone
kund_list = []  # Kundobjekt lagras här

outputDocument = Document(template) # Definiera en output-fil från Word-mallen


def duplicate_content_on_new_page(document, template):
    document.add_page_break()

    for index, paragraph in enumerate(template.paragraphs):
        new_paragraph = document.add_paragraph()
        new_paragraph.style = paragraph.style

        # Apply 9 cm indentation to the first 11 lines
        if index < 11:
            new_paragraph.paragraph_format.left_indent = Inches(3.6)

        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)

            # Copy formatting from the original run to the new run
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = run.font.size
            new_run.font.name = run.font.name
            if run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.highlight_color:
                new_run.font.highlight_color = run.font.highlight_color

with open(kunddata, mode='r', encoding='utf-8') as csvfile:
    csvreader = csv.reader(csvfile, delimiter='\t')  # Use tab as delimiter
    for row in csvreader:
        kund = Kund(*row)  # Gör en kund av varje rad i CSV-filen
        kund.fornamn = kund.fornamn.title()
        kund.efternamn = kund.efternamn.title()
        kund.adress1 = kund.adress1.title()
        if "lgh" in kund.adress1.lower():
            kund.adress1 = kund.adress1.replace("Lgh", "lgh")
        kund.adress21 = kund.adress21.replace(" ", "")
        if len(kund.adress21) > 3:  
            kund.adress21 = kund.adress21[:3] + " " + kund.adress21[3:] # Lägg till mellanslag efter de första tre siffrorna
        kund.adress22 = kund.adress22.title()
            
        kund_list.append(kund)
        duplicate_content_on_new_page(outputDocument, Document(template))   
        
        for paragraph in outputDocument.paragraphs:  # Loop through all paragraphs in the document
            if "{namn}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{namn}", kund.fornamn + " " + kund.efternamn)
            if "{adress1}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{adress1}", kund.adress1)
            if "{adress21}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{adress21}", kund.adress21)
            if "{adress22}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{adress22}", kund.adress22)
            if "{today}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{today}", "Visby " + today)
                
# Spara det uppdaterade dokumentet
outputDocument.save(output_path)